from datetime import date

from docx import Document

from app.services.docx_fill import generate_docx_from_template
from app.services.mandat_tokens import build_mandat_mapping
from app.services.token_utils import extract_docx_tokens


def test_mandat_md_signature_tokens_present(tmp_path):
    doc = Document()
    doc.add_paragraph("Signature jour : «MANDAT_JOUR_SIGNATURE»")
    doc.add_paragraph("Signature mois/année : «MANDAT_DATE_SIGNATURE»")
    tpl_path = tmp_path / "mandat_md.docx"
    doc.save(tpl_path)

    mapping = build_mandat_mapping({}, date(2025, 5, 12))
    out_path = tmp_path / "out.docx"

    report = generate_docx_from_template(str(tpl_path), str(out_path), mapping, strict=True)

    assert report.ok
    assert extract_docx_tokens(str(out_path)) == set()
