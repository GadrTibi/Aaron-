from docx import Document

from app.services.token_utils import extract_docx_tokens


def test_docx_token_extraction_handles_split_runs(tmp_path):
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("«TO")
    p.add_run("KEN»")
    doc_path = tmp_path / "split_runs.docx"
    doc.save(doc_path)

    tokens = extract_docx_tokens(str(doc_path))

    assert tokens == {"«TOKEN»"}
