from datetime import date
from pathlib import Path

from docx import Document

from app.services.docx_fill import generate_docx_from_template
from app.services.mandat_tokens import build_mandat_mapping


def test_signature_date_injected_when_tokens_absent(tmp_path: Path):
    template_path = tmp_path / "mandat_cd.docx"
    output_path = tmp_path / "out.docx"

    doc = Document()
    doc.add_paragraph("Fait à Paris, le ")
    doc.save(template_path)

    sig_date = date(2025, 12, 30)
    mapping = build_mandat_mapping({}, signature_date=sig_date)

    generate_docx_from_template(str(template_path), str(output_path), mapping)

    filled = Document(output_path)
    texts = [p.text for p in filled.paragraphs]
    assert any("Fait à Paris, le 30 décembre 2025" in t for t in texts)
