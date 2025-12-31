import os
from typing import Dict

from docx import Document

from app.services.generation_report import GenerationReport
from app.services.token_utils import extract_docx_tokens_from_document


def _replace_in_paragraph(paragraph, mapping: Dict[str, str]) -> None:
    def rebuild():
        segs = [(i, r.text or "") for i, r in enumerate(paragraph.runs)]
        joined = "".join(t for _, t in segs)
        return segs, joined

    for key, val in mapping.items():
        while True:
            segs, text = rebuild()
            idx = text.find(key)
            if idx < 0:
                break
            pos = 0
            start_run = start_off = end_run = end_off = 0
            for i, t in segs:
                if idx >= pos and idx <= pos + len(t):
                    start_run = i
                    start_off = idx - pos
                    break
                pos += len(t)
            pos = 0
            end_idx = idx + len(key)
            for i, t in segs:
                if end_idx > pos and end_idx <= pos + len(t):
                    end_run = i
                    end_off = end_idx - pos
                    break
                pos += len(t)
            sr = paragraph.runs[start_run]
            original_text = sr.text
            pre = original_text[:start_off]
            suffix = original_text[end_off:]
            sr.text = pre + val
            for ridx in range(start_run + 1, end_run):
                paragraph.runs[ridx].text = ""
            if end_run != start_run:
                lr = paragraph.runs[end_run]
                suffix = lr.text[end_off:]
                lr.text = suffix
            else:
                sr.text = sr.text + suffix


def _replace_in_document(doc, mapping: Dict[str, str]) -> None:
    for p in doc.paragraphs:
        _replace_in_paragraph(p, mapping)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p, mapping)


def replace_placeholders_docx(template_path: str, output_path: str, mapping: Dict[str, str]) -> None:
    doc = Document(template_path)
    _replace_in_document(doc, mapping)
    doc.save(output_path)


def _collect_leftovers(doc) -> set[str]:
    return extract_docx_tokens_from_document(doc)


def inject_signature_date_if_missing(doc: Document, date_full_str: str, report: GenerationReport | None = None) -> bool:
    """Injecte la date sur la ligne «Fait à Paris, le» si le template n'a pas les tokens MD.

    Retourne True si le document a été modifié.
    """

    modified = False
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if text.strip() != "Fait à Paris, le":
            continue
        if any(ch.isdigit() for ch in text) or "MANDAT_" in text:
            continue
        paragraph.text = f"Fait à Paris, le {date_full_str}"
        modified = True
        if report:
            report.add_note("Date de signature injectée sur la ligne 'Fait à Paris, le' (template sans tokens).")
        break
    return modified


def generate_docx_from_template(
    template_path: str,
    output_path: str,
    mapping: Dict[str, str],
    *,
    strict: bool = False,
) -> GenerationReport:
    """Génère un DOCX en remplaçant les tokens et signale ceux restants.

    Retourne un GenerationReport pour afficher les warnings dans l'UI.
    """
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template DOCX introuvable: {template_path}")
    doc = Document(template_path)
    _replace_in_document(doc, mapping)

    report = GenerationReport()
    date_full = mapping.get("MANDAT_DATE_SIGNATURE_FULL")
    if isinstance(date_full, str) and date_full.strip():
        inject_signature_date_if_missing(doc, date_full, report)

    leftovers = _collect_leftovers(doc)

    doc.save(output_path)
    if leftovers:
        report.add_missing_tokens(sorted(leftovers), blocking=strict)
        report.add_note("Certains tokens n'ont pas été remplacés dans le DOCX.")
    return report
